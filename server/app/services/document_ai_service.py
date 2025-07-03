from typing import List
import os
from pathlib import Path
from google.api_core.client_options import ClientOptions
from google.cloud import documentai
from langchain.schema import Document as LangChainDocument

from app.core.config import settings


class DocumentAIService:
    """
    Service để xử lý documents sử dụng Google Cloud Document AI
    thay thế cho DocLing
    """

    # Giới hạn kích thước file cho Document AI (40MB để an toàn)
    MAX_FILE_SIZE_BYTES = 40 * 1024 * 1024  # 40MB

    def __init__(self):
        """
        Khởi tạo Document AI client
        """
        # Set api_endpoint nếu sử dụng location khác "us"
        if settings.DOCUMENT_AI_LOCATION != "us":
            opts = ClientOptions(
                api_endpoint=f"{settings.DOCUMENT_AI_LOCATION}-documentai.googleapis.com"
            )
            self.client = documentai.DocumentProcessorServiceClient(client_options=opts)
        else:
            self.client = documentai.DocumentProcessorServiceClient()

        # Tạo processor name
        self.processor_name = self.client.processor_path(
            settings.GOOGLE_CLOUD_PROJECT_ID,
            settings.DOCUMENT_AI_LOCATION,
            settings.DOCUMENT_AI_PROCESSOR_ID,
        )

    def load_documents(self, file_path: str) -> List[LangChainDocument]:
        """
        Load và xử lý document từ file path

        Args:
            file_path (str): Đường dẫn đến file cần xử lý

        Returns:
            List[LangChainDocument]: Danh sách documents đã được xử lý

        Raises:
            Exception: Nếu có lỗi trong quá trình xử lý
        """
        try:
            # Kiểm tra kích thước file trước khi xử lý
            file_size = os.path.getsize(file_path)
            if file_size > self.MAX_FILE_SIZE_BYTES:
                return self._process_large_file_fallback(file_path, file_size)

            # Đọc file content
            with open(file_path, "rb") as file:
                file_content = file.read()

            # Xác định MIME type dựa trên file extension
            file_extension = Path(file_path).suffix.lower()
            mime_type = self._get_mime_type(file_extension)

            # Tạo raw document
            raw_document = documentai.RawDocument(
                content=file_content, mime_type=mime_type
            )

            # Tạo process request
            request = documentai.ProcessRequest(
                name=self.processor_name, raw_document=raw_document
            )

            # Xử lý document
            result = self.client.process_document(request=request)
            document = result.document

            # Chuyển đổi thành LangChain Document format
            langchain_documents = self._convert_to_langchain_documents(
                document, file_path
            )

            return langchain_documents

        except Exception as e:
            raise Exception(f"Lỗi khi xử lý document với Document AI: {str(e)}")

    def _process_large_file_fallback(
        self, file_path: str, file_size: int
    ) -> List[LangChainDocument]:
        """
        Xử lý file lớn vượt quá giới hạn Document AI bằng fallback methods

        Args:
            file_path (str): Đường dẫn đến file
            file_size (int): Kích thước file

        Returns:
            List[LangChainDocument]: Documents đã được xử lý

        Raises:
            Exception: Nếu không thể xử lý file
        """
        file_extension = Path(file_path).suffix.lower()
        file_size_mb = file_size / (1024 * 1024)

        print(
            f"File quá lớn ({file_size_mb:.1f}MB > {self.MAX_FILE_SIZE_BYTES / (1024 * 1024)}MB), sử dụng fallback method..."
        )

        try:
            if file_extension == ".pdf":
                return self._process_pdf_fallback(file_path)
            elif file_extension in [".txt"]:
                return self._process_text_fallback(file_path)
            elif file_extension in [".docx"]:
                return self._process_docx_fallback(file_path)
            else:
                raise Exception(
                    f"File quá lớn ({file_size_mb:.1f}MB) và không hỗ trợ fallback method cho định dạng {file_extension}. "
                    f"Vui lòng sử dụng file nhỏ hơn {self.MAX_FILE_SIZE_BYTES / (1024 * 1024)}MB hoặc chia nhỏ file."
                )
        except Exception as e:
            raise Exception(f"Không thể xử lý file lớn: {str(e)}")

    def _process_pdf_fallback(self, file_path: str) -> List[LangChainDocument]:
        """
        Xử lý PDF file lớn bằng PyPDF2

        Args:
            file_path (str): Đường dẫn đến PDF file

        Returns:
            List[LangChainDocument]: Documents đã được xử lý
        """
        try:
            import PyPDF2

            documents = []
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Tạo document tổng hợp
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    full_text += page_text + "\n"

                    # Tạo document cho từng page
                    if page_text.strip():
                        page_doc = LangChainDocument(
                            page_content=page_text,
                            metadata={
                                "source": os.path.basename(file_path),
                                "source_type": "pdf_fallback",
                                "page_number": page_num + 1,
                                "total_pages": len(pdf_reader.pages),
                                "processor": "pypdf2_fallback",
                            },
                        )
                        documents.append(page_doc)

                # Tạo document tổng hợp
                if full_text.strip():
                    main_doc = LangChainDocument(
                        page_content=full_text,
                        metadata={
                            "source": os.path.basename(file_path),
                            "source_type": "pdf_fallback",
                            "total_pages": len(pdf_reader.pages),
                            "processor": "pypdf2_fallback",
                        },
                    )
                    documents.insert(0, main_doc)  # Thêm vào đầu list

                return documents

        except ImportError:
            raise Exception(
                "PyPDF2 chưa được cài đặt. Vui lòng cài đặt: pip install PyPDF2"
            )
        except Exception as e:
            raise Exception(f"Lỗi khi xử lý PDF với fallback method: {str(e)}")

    def _process_text_fallback(self, file_path: str) -> List[LangChainDocument]:
        """
        Xử lý text file lớn

        Args:
            file_path (str): Đường dẫn đến text file

        Returns:
            List[LangChainDocument]: Documents đã được xử lý
        """
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()

            document = LangChainDocument(
                page_content=content,
                metadata={
                    "source": os.path.basename(file_path),
                    "source_type": "text_fallback",
                    "processor": "text_fallback",
                },
            )

            return [document]

        except Exception as e:
            raise Exception(f"Lỗi khi xử lý text file với fallback method: {str(e)}")

    def _process_docx_fallback(self, file_path: str) -> List[LangChainDocument]:
        """
        Xử lý DOCX file lớn bằng python-docx

        Args:
            file_path (str): Đường dẫn đến DOCX file

        Returns:
            List[LangChainDocument]: Documents đã được xử lý
        """
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            full_text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"

            document = LangChainDocument(
                page_content=full_text,
                metadata={
                    "source": os.path.basename(file_path),
                    "source_type": "docx_fallback",
                    "processor": "python-docx_fallback",
                },
            )

            return [document]

        except ImportError:
            raise Exception(
                "python-docx chưa được cài đặt. Vui lòng cài đặt: pip install python-docx"
            )
        except Exception as e:
            raise Exception(f"Lỗi khi xử lý DOCX với fallback method: {str(e)}")

    def _get_mime_type(self, file_extension: str) -> str:
        """
        Xác định MIME type dựa trên file extension

        Args:
            file_extension (str): Phần mở rộng của file

        Returns:
            str: MIME type tương ứng
        """
        mime_types = {
            ".pdf": "application/pdf",
            ".txt": "text/plain",
            ".doc": "application/msword",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".tiff": "image/tiff",
            ".tif": "image/tiff",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
        }

        return mime_types.get(file_extension, "application/octet-stream")

    def _convert_to_langchain_documents(
        self, document: documentai.Document, source_file: str
    ) -> List[LangChainDocument]:
        """
        Chuyển đổi Document AI result thành LangChain Documents

        Args:
            document: Document AI processed document
            source_file (str): Đường dẫn file gốc

        Returns:
            List[LangChainDocument]: Danh sách LangChain documents
        """
        documents = []

        # Nếu có text content, tạo document chính
        if document.text:
            main_doc = LangChainDocument(
                page_content=document.text,
                metadata={
                    "source": os.path.basename(source_file),
                    "source_type": "document_ai",
                    "total_pages": len(document.pages) if document.pages else 1,
                    "processor": "google_document_ai",
                },
            )
            documents.append(main_doc)

        # Xử lý từng page nếu có thông tin chi tiết
        for page_idx, page in enumerate(document.pages):
            # Lấy text của page
            page_text = self._extract_page_text(document, page)

            if page_text.strip():  # Chỉ tạo document nếu có content
                page_doc = LangChainDocument(
                    page_content=page_text,
                    metadata={
                        "source": os.path.basename(source_file),
                        "source_type": "document_ai",
                        "page_number": page_idx + 1,
                        "total_pages": len(document.pages),
                        "page_width": page.dimension.width if page.dimension else None,
                        "page_height": (
                            page.dimension.height if page.dimension else None
                        ),
                        "processor": "google_document_ai",
                    },
                )
                documents.append(page_doc)

        # Nếu không có page-level content, chỉ trả về document chính
        if len(documents) == 1:
            return documents

        # Trả về cả document chính và page-level documents
        return documents

    def _extract_page_text(
        self, document: documentai.Document, page: documentai.Document.Page
    ) -> str:
        """
        Trích xuất text từ một page cụ thể

        Args:
            document: Document AI processed document
            page: Page object từ document

        Returns:
            str: Text content của page
        """
        try:
            # Lấy text từ các text segments trong page
            page_text = ""

            # Nếu có tokens, sử dụng tokens để xây dựng text
            if hasattr(page, "tokens") and page.tokens:
                for token in page.tokens:
                    if (
                        hasattr(token, "layout")
                        and token.layout
                        and hasattr(token.layout, "text_anchor")
                    ):
                        text_anchor = token.layout.text_anchor
                        if text_anchor and hasattr(text_anchor, "text_segments"):
                            for segment in text_anchor.text_segments:
                                start_index = (
                                    segment.start_index
                                    if hasattr(segment, "start_index")
                                    else 0
                                )
                                end_index = (
                                    segment.end_index
                                    if hasattr(segment, "end_index")
                                    else len(document.text)
                                )
                                page_text += document.text[start_index:end_index]

            # Nếu không có tokens, fallback về paragraphs
            elif hasattr(page, "paragraphs") and page.paragraphs:
                for paragraph in page.paragraphs:
                    if (
                        hasattr(paragraph, "layout")
                        and paragraph.layout
                        and hasattr(paragraph.layout, "text_anchor")
                    ):
                        text_anchor = paragraph.layout.text_anchor
                        if text_anchor and hasattr(text_anchor, "text_segments"):
                            for segment in text_anchor.text_segments:
                                start_index = (
                                    segment.start_index
                                    if hasattr(segment, "start_index")
                                    else 0
                                )
                                end_index = (
                                    segment.end_index
                                    if hasattr(segment, "end_index")
                                    else len(document.text)
                                )
                                page_text += document.text[start_index:end_index] + "\n"

            return page_text.strip()

        except Exception as e:
            # Nếu có lỗi, trả về empty string
            print(f"Warning: Không thể trích xuất text từ page: {str(e)}")
            return ""


def get_document_ai_service():
    return DocumentAIService()
